'''
    Title:  Automated Change Mangement Agenda 0.8.1
    Author:  Joe Friedrich
    License:  MIT
'''

import csv
import datetime
import os
from docx import Document
from dateutil import tz
from tkinter import Tk
from tkinter.filedialog import askopenfilename

def open_agenda_template():
    '''
        Write!
    '''
    document = Document(r'Location of your GenericAgenda.docx goes here')
    return document

def get_input_file_name():
    '''
        Write!
    '''
    pause = input("Please hit the enter key when you are ready to select a CSV file.")
    window = Tk()
    window.withdraw()
    file_name = askopenfilename()
    return file_name

def get_data(file_name):
    '''
        Write!
    '''
    open_file = open(file_name, encoding='ansi')
    data = list(csv.reader(open_file))
    open_file.close()
    return data

def get_user_date():
    '''
        Write!
    '''
    user_date = input('What is the date of the next change management meeting? (mm/dd/yy)  ')
    return user_date
    
def get_tz_database():
    '''
        Write!
    '''
    open_file = open(r'Location of the timeZone.CSV [lotus notes to windows timezone conversion file] goes here.')
    data = dict(csv.reader(open_file))
    open_file.close()
    return data
    
def calculate_timezone(start_time, end_time, local_timezone, standard_timezone):
    '''
        Write!
    '''
    ocr_naive_start = datetime.datetime.strptime(start_time, "%m/%d/20%y %I - %M %p")
    ocr_naive_end = datetime.datetime.strptime(end_time, "%m/%d/20%y %I - %M %p")
    ocr_timezone = tz.tzwin(local_timezone)
    ocr_aware_start = ocr_naive_start.replace(tzinfo=ocr_timezone)
    ocr_aware_end = ocr_naive_end.replace(tzinfo=ocr_timezone)
    standardized_tz_start = ocr_aware_start.astimezone(standard_timezone)
    standardized_tz_end = ocr_aware_end.astimezone(standard_timezone)
    return standardized_tz_start, standardized_tz_end

def fill_agenda_tables(agenda, ocr_data):
    '''
        Write!
    '''
    tables = agenda.tables #creates list of tables in the agenda doc
    for table, ocrs in zip(tables, ocr_data): #zip means this will stop when first item runs out of elements
        if len(ocrs) > 0: #if the ocr section is not blank
            first_row_data = ocrs.pop(0)
            first_row_table = table.rows[1].cells
            for item, cell in zip(first_row_data, first_row_table):
                cell.text = item
        if len(ocrs) > 0: #if there is at least a second ocr, add a row and the data
            for ocr in ocrs:
                row = table.add_row().cells
                for element, box in zip(ocr, row):
                    box.text = element
    return agenda
    
def get_docx_output_location():
    '''
        Write!
    '''
    user_folder = r'Location of the'
    today = datetime.datetime.today()
    today_text = today.strftime('%m-%d-%y ')
    filename = user_folder + today_text + 'Agenda.docx'
    return filename
    
#---------------------Begin Program--------------------
agenda = open_agenda_template()
unified_timezone = tz.tzwin('Eastern Standard Time') #agreed upon standard timezone
tz_database = get_tz_database()

data = get_data(get_input_file_name())
garbage = data.pop(0) #removes headers
garbage = data.pop(0) #removes headers

user_date = get_user_date()

naive_meeting_date_time = datetime.datetime.strptime(user_date + " 9:30 AM", "%m/%d/%y %I:%M %p")
meeting_date_time = naive_meeting_date_time.replace(tzinfo=unified_timezone)
meeting_date = meeting_date_time.strftime('%b %d, %Y')
print("\ncurrent meeting date/time " + meeting_date)

previous_week = datetime.timedelta(days=7,hours=1)
previous_meeting_time = meeting_date_time - previous_week #one hour before the previous meeting
previous_meeting_date = previous_meeting_time.strftime('%b %d, %Y')
print("previous meeting date/time " + previous_meeting_date)

next_week = datetime.timedelta(days=7)
next_meeting_time = meeting_date_time + next_week #the next meeting
next_meeting_date = next_meeting_time.strftime('%b %d, %Y')
print("next meeting date/time " + next_meeting_date)

agenda.paragraphs[2].text = meeting_date
agenda.paragraphs[32].text = 'Emergency changes logged ' + previous_meeting_date + ' - ' + meeting_date + ' 8:30am'
agenda.paragraphs[39].text = previous_meeting_date + ' - ' + meeting_date

emergency_ocr = []
low_pre_approved = []
pre_approved = []
low = []
medium = []
high = []
open_5_days = []
previously_approved_changes = []

for ocr in data:
    emergency = ocr[7]
    priority = ocr[6]
    status = ocr[5]
    
    local_timezone = tz_database[ocr[3]]
    
    unified_tz_start, unified_tz_end = calculate_timezone(ocr[0], 
                                                        ocr[1], 
                                                        local_timezone, 
                                                        unified_timezone)
    
    ocr = [unified_tz_start.strftime('%a %b %d %H:%M %Y'),
            unified_tz_end.strftime('%a %b %d %H:%M %Y'),
            ocr[4],
            ocr[2],
            ocr[16]]

    if emergency == 'Yes':
        if unified_tz_start > previous_meeting_time:
            emergency_ocr.append(ocr)
    elif status == "Awaiting Opp Approval":
        if unified_tz_start < meeting_date_time:
                if priority[0] == "3":
                    low_pre_approved.append(ocr)
                else:
                    pre_approved.append(ocr)
        elif priority[0] == "3":
            low.append(ocr)
        elif priority[0] == "2":
            medium.append(ocr)
        else:
            high.append(ocr)
    elif status == "Approved/Scheduled":
        if unified_tz_end < previous_meeting_time:
            open_5_days.append(ocr)
        elif meeting_date_time < unified_tz_start < next_meeting_time:
            previously_approved_changes.append(ocr)

agenda_data = [previously_approved_changes,
               low_pre_approved,
               low,
               open_5_days,
               emergency_ocr,
               pre_approved,
               high,
               medium]

agenda = fill_agenda_tables(agenda, agenda_data)

agenda.save(get_docx_output_location())
